from flask import Flask, render_template, request, send_file, redirect, url_for, flash
import io, datetime, re, csv, zipfile
from fugashi import Tagger
from docx import Document as DocxDocument

app = Flask(__name__)
app.secret_key = "replace-this"

# ---- 漢字判定 ----
KANJI_CHAR_RE = re.compile(r'[一-鿿]')

# ---- 学年CSV 読み込み ----
def load_kanji_grade_mapping(csv_path='kanji_grade.csv'):
    mapping = {}
    try:
        with open(csv_path, encoding='utf-8') as f:
            reader = csv.reader(f)
            for row in reader:
                if len(row) < 2:
                    continue
                kanji = row[0].strip()
                try:
                    grade = int(row[1])
                except ValueError:
                    continue
                if kanji and 1 <= grade <= 9:
                    mapping[kanji] = grade
    except FileNotFoundError:
        print(f"[Warning] {csv_path} が見つかりません。すべての漢字にフリガナを振ります。")
    return mapping

KANJI_GRADE = load_kanji_grade_mapping()

# ---- MeCab(fugashi) ----
tagger = Tagger()

# ---- 改行保持版 ルビ付与（テキスト→テキスト）----
def annotate_by_grade(text: str, threshold: int) -> str:
    out_chunks = []
    for chunk in re.split(r'(\r\n|\r|\n)', text):
        if chunk in ('\r\n', '\r', '\n'):
            out_chunks.append(chunk)
            continue

        processed = []
        if not chunk:
            out_chunks.append(chunk)
            continue

        for token in tagger(chunk):
            surface = token.surface

            # feature末尾あたりにあるカタカナ読みを拾う
            reading = None
            for feat in reversed(token.feature):
                if feat and re.fullmatch(r'[\u30A0-\u30FF]+', feat):
                    reading = feat
                    break
            if reading:
                try:
                    import jaconv
                    reading = jaconv.kata2hira(reading)
                except ImportError:
                    pass

            if not reading or not KANJI_CHAR_RE.search(surface):
                processed.append(surface)
                continue

            m = re.match(r'^([一-鿿]+)(.*)$', surface)
            if m:
                kanji_part, rest = m.group(1), m.group(2)
            else:
                kanji_part, rest = surface, ''

            if rest and len(kanji_part) == 1:
                kanji_reading = reading[:-len(rest)] if len(reading) > len(rest) else reading
            else:
                kanji_reading = reading

            grades = [KANJI_GRADE.get(ch) for ch in kanji_part]
            if any(g is None for g in grades):
                attach = True
            elif all(g < threshold for g in grades):
                attach = False
            else:
                attach = True

            if not attach:
                processed.append(surface)
            else:
                processed.append(f"{kanji_part}（{kanji_reading}）{rest}")

        out_chunks.append(''.join(processed))

    return ''.join(out_chunks)

# ---- .docx構造を保ったままルビ挿入（段落＋表セル）----
def annotate_docx_inplace(doc: DocxDocument, threshold: int) -> None:
    # 段落
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                run.text = annotate_by_grade(run.text, threshold)

    # 表（セル内の段落／入れ子の表にも対応）
    def walk_tables(tables):
        for tbl in tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text:
                                run.text = annotate_by_grade(run.text, threshold)
                    # 入れ子テーブルがあれば再帰
                    if cell.tables:
                        walk_tables(cell.tables)
    walk_tables(doc.tables)

# ---- template.docm のマクロを保ちつつ document.xml を差し替え ----
def make_docm_from_xml_bytes(new_document_xml: bytes) -> io.BytesIO:
    template_path = 'template.docm'
    with zipfile.ZipFile(template_path, 'r') as zin:
        parts = {name: zin.read(name) for name in zin.namelist()}
    parts['word/document.xml'] = new_document_xml
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w') as zout:
        for name, data in parts.items():
            zout.writestr(name, data)
    buf.seek(0)
    return buf

# ---- プレーンテキスト入力 → docx作成 → XML抽出 ----
def text_to_document_xml(text: str) -> bytes:
    tmp = io.BytesIO()
    d = DocxDocument()
    for line in text.splitlines():
        d.add_paragraph(line)
    d.save(tmp)
    tmp.seek(0)
    with zipfile.ZipFile(tmp, 'r') as z:
        return z.read('word/document.xml')

# ---- .docx入力 → そのまま構造走査で注釈 → XML抽出 ----
def docx_bytes_to_document_xml(docx_bytes: bytes, threshold: int) -> bytes:
    stream = io.BytesIO(docx_bytes)
    doc = DocxDocument(stream)
    annotate_docx_inplace(doc, threshold)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    with zipfile.ZipFile(out, 'r') as z:
        return z.read('word/document.xml')

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        grade_str = request.form.get('grade', '').strip()
        if not grade_str.isdecimal():
            flash('学年を正しく選択してください', 'danger')
            return redirect(url_for('index'))
        threshold_grade = int(grade_str)

        uploaded = request.files.get('text_file')
        if uploaded and uploaded.filename:
            filename = uploaded.filename.lower()
            data = uploaded.read()

            # ---- .docx は構造保持で処理（表を含めて維持）----
            if filename.endswith('.docx'):
                new_xml = docx_bytes_to_document_xml(data, threshold_grade)
                docm_io = make_docm_from_xml_bytes(new_xml)
                filename_out = f"annotated_{datetime.datetime.now():%Y%m%d_%H%M%S}.docm"
                return send_file(
                    docm_io,
                    as_attachment=True,
                    download_name=filename_out,
                    mimetype='application/vnd.ms-word.document.macroEnabled.12'
                )
            # ---- .txt ほかプレーンテキストは従来処理 ----
            else:
                raw = ''
                for enc in ('utf-8', 'cp932'):
                    try:
                        raw = data.decode(enc)
                        break
                    except UnicodeDecodeError:
                        continue
                if not raw:
                    flash('テキストファイルが UTF-8 でも CP932 でもデコードできません', 'danger')
                    return redirect(url_for('index'))
                annotated = annotate_by_grade(raw, threshold_grade)
                new_xml = text_to_document_xml(annotated)
                docm_io = make_docm_from_xml_bytes(new_xml)
                filename_out = f"annotated_{datetime.datetime.now():%Y%m%d_%H%M%S}.docm"
                return send_file(
                    docm_io,
                    as_attachment=True,
                    download_name=filename_out,
                    mimetype='application/vnd.ms-word.document.macroEnabled.12'
                )
        else:
            # 画面のテキスト入力（先頭末尾の改行も保持）
            raw = request.form.get('source_text', '')
            if not raw:
                flash('ファイルをアップロードするか、テキストを入力してください', 'danger')
                return redirect(url_for('index'))
            annotated = annotate_by_grade(raw, threshold_grade)
            new_xml = text_to_document_xml(annotated)
            docm_io = make_docm_from_xml_bytes(new_xml)
            filename_out = f"annotated_{datetime.datetime.now():%Y%m%d_%H%M%S}.docm"
            return send_file(
                docm_io,
                as_attachment=True,
                download_name=filename_out,
                mimetype='application/vnd.ms-word.document.macroEnabled.12'
            )

    return render_template('index.html')

if __name__ == '__main__':
    app.run(debug=True, port=5000)
